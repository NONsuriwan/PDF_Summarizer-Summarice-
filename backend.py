import os
import PyPDF2
from transformers import T5ForConditionalGeneration, T5Tokenizer, AutoModel, AutoTokenizer
import torch
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from docx import Document
import nltk
from nltk.corpus import stopwords
import re

# Download necessary NLTK data
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger')
nltk.download('punkt_tab')

model_path = r'results_t5bbc_billsum_resume\checkpoint-130'
model = T5ForConditionalGeneration.from_pretrained(model_path)
tokenizer = T5Tokenizer.from_pretrained('results_t5bbc_billsum_resume')

# Load BERT model for similarity calculation
device = "cuda" if torch.cuda.is_available() else "cpu"
bert_model_name = "bert-base-uncased"
bert_tokenizer = AutoTokenizer.from_pretrained(bert_model_name)
bert_model = AutoModel.from_pretrained(bert_model_name).to(device)

def preprocess_text(text):
    text = text.lower()
    text = re.sub('[^a-zA-Z]', ' ', text)
    words = text.split()
    stop_words = set(stopwords.words("english"))
    filtered_words = [word for word in words if word not in stop_words]
    return " ".join(filtered_words)

def get_embeddings(text):
    inputs = bert_tokenizer(str(text), return_tensors="pt", truncation=True, padding=True).to(device)
    with torch.no_grad():
        outputs = bert_model(**inputs)
    embeddings = outputs.last_hidden_state.mean(dim=1).detach().to("cpu").numpy()
    return embeddings

def calculate_similarity(text1, text2):
    emb1 = get_embeddings(text1)
    emb2 = get_embeddings(text2)
    similarity = (cosine_similarity(emb1, emb2)[0][0]) * 100
    return similarity

def summarize_text(text):
    inputs = tokenizer.encode("summarize: " + text, return_tensors="pt", max_length=1024, truncation=True)
    outputs = model.generate(
        inputs, 
        max_length=300, 
        min_length=100, 
        length_penalty=1.2, 
        num_beams=5, 
        early_stopping=True
    )
    summary = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return summary

def summarize_pdf(file_path, job_description):
    with open(file_path, "rb") as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()

        # Preprocess and calculate similarity
        preprocessed_resume = preprocess_text(text)
        preprocessed_job_desc = preprocess_text(job_description)
        similarity = calculate_similarity(preprocessed_resume, preprocessed_job_desc)
        
        # Summarize the text
        summary = summarize_text(text)
        return summary, similarity

def process_pdfs_in_folder_bart(folder_path, output_docx_path, job_description, progress_callback=None):
    summaries_with_scores = []
    showgui = []
    pdf_files = [f for f in os.listdir(folder_path) if f.endswith('.pdf')]
    total_files = len(pdf_files)

    for i, file_name in enumerate(pdf_files):
        file_path = os.path.join(folder_path, file_name)
        print(f"Processing {file_name}...")
        
        # Process the PDF and summarize
        summary, similarity = summarize_pdf(file_path, job_description)
        summaries_with_scores.append((summary, similarity, file_name))
        
        # Send progress after each file is processed
        if progress_callback:
            progress_value = int((i + 1) / total_files * 100)
            progress_callback(progress_value)
            print(f"Progress: {progress_value}% to Front")

    # Sort summaries by similarity score
    summaries_with_scores.sort(key=lambda x: x[1], reverse=True)

    # Create DOCX file
    doc = Document()
    doc.add_heading(os.path.basename(folder_path), 0)

    for idx, (summary, similarity, file_name) in enumerate(summaries_with_scores):
        file_name_without_ext = os.path.splitext(file_name)[0]
        if idx < 5:
            showgui.append((file_name_without_ext, similarity))
        doc.add_heading(f'{file_name_without_ext} | Similarity: {similarity:.4f}', level=1)
        doc.add_paragraph(summary)

    # Save the DOCX file
    doc.save(output_docx_path)
    print(f"Output saved to {output_docx_path}")

    # Send final progress update after processing all files
    if progress_callback:
        progress_callback(100)
        print("Progress: 100% to Front")

    return showgui

# Example use case (you can remove this block in actual production)
if __name__ == "__main__":
    folder_path = "CHEF"
    output_docx_path = "summary.docx"
    job_description = "We are looking for a skilled chef with experience in Italian cuisine..."  # Example job description
    process_pdfs_in_folder_bart(folder_path, output_docx_path, job_description)